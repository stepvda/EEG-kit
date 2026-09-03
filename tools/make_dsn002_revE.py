#!/usr/bin/env python3
"""
make_dsn002_revE.py -- produce DSN-EEG-002 Rev E from the Rev D document.

Rev D is a figure-heavy document: sixteen rendered illustrations that took real work and
that this package still needs. Rather than re-author it and lose them, Rev E is made by
editing the Rev D file in place, so every figure survives and only the text that is wrong
changes.

What Rev E changes:
  * the revision block and the companion-document list
  * a new "What changed in Revision E" section recording the fourteen ECOs that touch this
    document, and the figure-label renaming
  * FIGURE LABELS: Rev D used HM-xx for both figures and parts, and the two namespaces
    collided -- HM-07 was the boom microphone arm in section 10 and the battery hatch in
    the STL set. Figures become FIG-nn; HM-xx is reserved for parts (ECO-EEG-015)
  * section 6: the harness is now two cables, not one twenty-way
  * section 7: the pod, superseded, now points at the released POD-P1 model with its
    panel openings
  * section 10: the parts table, corrected and extended
  * section 12: each objection gains its status in package v2
  * section 13: the schematic figures now point at SCH-EEG-005, and the numbers are the
    fitted ones

Licence: CC BY-SA 4.0.
"""
from __future__ import annotations
import copy
import os
import re
import sys

import docx
from docx.shared import Pt, RGBColor

SRC = ("/Users/nstephane/Dev/onevoice/documentation/Study - EEG/EEG-kit-RFQ/"
       "package/docs/DSN-EEG-002_RevD_design_and_assembly.docx")
DST = ("/Users/nstephane/Dev/onevoice/documentation/Study - EEG/EEG-kit-RFQ/"
       "package_v2.4/docs/DSN-EEG-002_RevE_design_and_assembly.docx")

# Rev D figure label -> Rev E figure label. HM-xx is reserved for parts from Rev E on.
FIGMAP = {
    "HM-00": "FIG-01", "HM-01": "FIG-02", "HM-02": "FIG-03", "HM-03": "FIG-04",
    "HM-14": "FIG-05", "HM-04": "FIG-06", "HM-05": "FIG-07", "HM-06": "FIG-08",
    "HM-15": "FIG-09", "HM-08": "FIG-10", "W1": "FIG-11", "POD-00": "FIG-12",
    "HM-09": "FIG-13", "HM-10": "FIG-14", "HM-11": "FIG-15", "HM-12": "FIG-16",
    "HM-13": "FIG-17", "HM-07": "FIG-18", "CASE-00": "FIG-19",
    "S1": "FIG-20", "S2": "FIG-21", "S3": "FIG-22", "S4": "FIG-23",
}


def set_text(p, text, bold=False, italic=False, size=None, color=None):
    for r in list(p.runs)[1:]:
        r._element.getparent().remove(r._element)
    if not p.runs:
        p.add_run("")
    r = p.runs[0]
    r.text = text
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)


def insert_after(p, text, style=None, bold=False, italic=False):
    new = copy.deepcopy(p._element)
    p._element.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        try:
            np_.style = style
        except Exception:
            pass
    set_text(np_, text, bold=bold, italic=italic)
    return np_


def main():
    d = docx.Document(SRC)
    P = d.paragraphs

    set_text(P[2], "Document: DSN-EEG-002   Revision: E   Date: 1 September 2026", bold=True)
    set_text(P[4],
             "Governing documents: DSN-EEG-003 Rev B (manufacturing design package), then "
             "RFQ-EEG-001 Rev D (requirements and acceptance), then ICD-EEG-006 Rev A "
             "(module interfaces). Part identifiers are governed by PARTS-EEG-019 Rev A. "
             "Where this document and tools/design.py disagree, design.py governs.")
    set_text(P[5], "Revision E — what changed and why", bold=True)
    set_text(P[6],
             "Rev E changes no illustration. It corrects the text where package v2 found it "
             "to be wrong, and it renames the figures. Rev D used the HM-xx prefix for both "
             "figure labels and part numbers, and the two namespaces collided: HM-07 was the "
             "boom microphone arm in section 10 and the battery hatch in DSN-EEG-003 section "
             "4, in the STL set, in the kit BOM and in the RFQ scope line. A manufacturer "
             "printing “HM-07” produced a different part depending on which document "
             "was open, and section 10 requires the part identifier to be engraved in the "
             "model, so the wrong identifier would have been moulded into the part. From "
             "Rev E, figures are FIG-nn and HM-xx means a part and nothing else. The battery "
             "hatch is HM-08. PARTS-EEG-019 is the single register and carries the full "
             "migration table.")

    anchor = P[6]
    for line in [
        "The other corrections in Rev E, each of them an ECO in ECO-EEG-016:",
        "• Section 6, wiring. The harness is now TWO cables, not one twenty-way. A "
        "12-way screened electrode bundle carries the eight scalp sites, the two ear "
        "references, the bias lead and the screen; a separate 10-way ribbon carries the "
        "eight contact-light lines, the light common and its return. Rev D ran all of it "
        "through one connector at the far edge of the analogue zone, which forced eight "
        "digital conductors across the whole front end and put them in the same cable as "
        "eight high-impedance electrode leads. DSN-EEG-003 Rev A.2 recorded that and "
        "accepted it; that was the wrong call (ECO-EEG-014). WH-EEG-008 is the wire list.",
        "• Section 5, contact lights. Each site carries a two-lead bicolour LED between "
        "its own line and a common phase line. Rev A of the carrier had no driver for them "
        "at all: the shift-register module's outputs were never brought to the board "
        "(ECO-EEG-001). They are dark at power-on because the phase line is an input at "
        "reset, so no current can flow whatever the register contains.",
        "• Section 7, the pod. POD-P1 is now a released parametric model with the panel "
        "openings of RFQ M-02 cut in it: three touch-proof DIN 42802 sockets, the headphone "
        "jack, the boom connector, the room-microphone port, a data USB-C on the isolator "
        "module and a SEPARATE charge-only USB-C. Rev D implied one connector for both, "
        "which cannot be reconciled with the isolation requirement (ECO-EEG-003).",
        "• Section 10, parts. The table is corrected and extended: the battery hatch is "
        "HM-08, the module plate MP-01 is new, and every part now names the file that "
        "defines it. Parts that had no file in Rev D — the TPU pads, the yoke, the chin "
        "strap, the service key, the fit-test coupon, the bottom foam layer — are "
        "listed with their status.",
        "• Section 11, the travel case. The internal size is about 340 × 250 × "
        "210 mm, which is what an assembled helmet standing upright needs. RFQ Rev C gave "
        "300 × 220 × 110 mm in one place and the larger figure in the kit BOM; "
        "the larger one is correct and RFQ Rev D says so.",
        "• Section 13, schematics. The figures here are the design intent. The released "
        "schematic is SCH-EEG-005 Rev B, eight sheets, generated from the same source as "
        "the board. Two circuits changed: the envelope detector is a full-wave rectifier "
        "into a 48.8 Hz Butterworth filter on a quad op-amp, because the Rev A dual left "
        "the filter out of circuit entirely (ECO-EEG-004 and 006), and the input protection "
        "is on the carrier EEG-CAR-01, not on a separate panel board.",
        "• Section 12, the adversarial review. Every objection now carries its status in "
        "package v2. Three of them are answered; the rest still are not, and say so.",
    ]:
        anchor = insert_after(anchor, line)

    # ---- figure labels -------------------------------------------------------------
    pat = re.compile(r"^(" + "|".join(sorted(FIGMAP, key=len, reverse=True)) + r")(\s*—|\s*--|\s*—)")
    renamed = 0
    for p in d.paragraphs:
        t = p.text
        m = pat.match(t.strip())
        if m and p.runs:
            old = m.group(1)
            new = FIGMAP[old]
            first = p.runs[0]
            if first.text.strip().startswith(old):
                first.text = first.text.replace(old, new, 1)
                renamed += 1

    # ---- section 6 harness caption and table ----------------------------------------
    for p in d.paragraphs:
        if p.text.strip().startswith("FIG-11") or p.text.strip().startswith("W1 —"):
            set_text(p, "FIG-11 — the harness as a schedule. Superseded in Rev E: the "
                        "harness is two cables. Eight scalp sites, two references and the "
                        "bias lead converge on a 12-way screened bundle; the eight contact "
                        "lights leave on a separate 10-way ribbon. WH-EEG-008 is the "
                        "controlled wire list.", italic=True)
    t = d.tables[7]
    fixes = {
        "Eight scalp sites": ("8 × screened", "Assembly base → channel in the arch "
                              "or halo → crown → rearward → pod, on the 12-way "
                              "screened bundle to J14",
                              "One conductor per site plus one overall screen. The screen is "
                              "grounded at the pod end only, at R91, so the frame carries no "
                              "loop."),
        "Reference": ("2 × screened", "Ear clip → temple → halo channel "
                      "→ pod, on the same 12-way bundle",
                      "Linked earlobes into SRB1, shared by all sixteen channels"),
        "Bias": ("1 × screened", "Fpz pad → halo front channel → pod, on the "
                 "same 12-way bundle",
                 "Driven common-mode return, not an earth. Leaves the carrier through R11."),
    }
    for row in t.rows[1:]:
        key = row.cells[0].text.strip()
        if key in fixes:
            for i, v in enumerate(fixes[key], start=1):
                row.cells[i].text = v
    # add the light cable as its own run
    new_row = t.add_row()
    new_row.cells[0].text = "Contact lights"
    new_row.cells[1].text = "10-way ribbon"
    new_row.cells[2].text = ("Eight assemblies → crown → rearward → pod, on a "
                             "separate 10-way ribbon to J30")
    new_row.cells[3].text = ("NEW IN REV E (ECO-EEG-014). Eight drive lines, the LED_V phase "
                             "common and its return. Kept out of the electrode bundle so no "
                             "digital conductor shares a cable with an electrode lead.")

    # ---- section 10 parts table -------------------------------------------------------
    t = d.tables[9]
    updates = {
        "HM-01": ("Frame monocoque: halo, sagittal arch, coronal arch, rail stubs, pod shell, "
                  "internal channels", "1", "PA12, MJF",
                  "One print, about 240 g. mech/stl/HM-01_frame_monocoque.stl. A STEP model "
                  "follows the Stage 0 fit measurement (section 12 objection 3)."),
        "HM-02": ("TPU pads: brow, occiput ×2, crown", "4 + 4 spare", "TPU 85A",
                  "Consumable, replaced each turnaround. "
                  "mech/step/HM-02_brow_pad.step — the other three follow the same form."),
        "HM-03": ("Occipital yoke and ratchet dial", "1", "PA12 + POM pawl",
                  "2 mm per click, 52–62 cm. Bought-in ratchet; the yoke model is a "
                  "Phase 1 deliverable (PARTS-EEG-019)."),
        "HM-04": ("Electrode assembly: body, spring, cup seat, bicolour light, gel port",
                  "8 + 2 spare", "PA12 + steel spring + bicolour LED",
                  "Bonded into the frame at manufacture. "
                  "mech/step/HM-04_electrode_assembly_body.step, drawing MECH-EEG-020."),
        "HM-05": ("Ag/AgCl cups on service bayonet", "8 + 2 spare", "Sintered Ag/AgCl",
                  "Bought in. Released only with the HM-09 service key. Replaced outright "
                  "every ~25 sessions (SVC-EEG-013)."),
        "HM-06": ("Chin strap and chin cup with liner", "1", "PA12 + webbing",
                  "Removable, and the runner says so before it is first mentioned. Liner "
                  "consumable."),
        "HM-07": ("Boom microphone arm", "1", "PA12 + gooseneck",
                  "Detaches at the temple. THIS IS HM-07 — the battery hatch is HM-08. "
                  "See PARTS-EEG-019 and ECO-EEG-015."),
        "HM-08": ("Battery hatch, quarter-turn, and keyed cell carrier", "1 each", "PA12",
                  "RENAMED IN REV E from HM-07. Tool-free, interlocked. "
                  "mech/step/HM-08_battery_hatch.step."),
        "HM-09": ("Service key for cup release", "1 per operator, not in the kit", "PA12",
                  "Deliberately absent from the participant's kit. Controlled item: "
                  "SVC-EEG-013 section 3. mech/step/HM-09_service_key.step."),
        "CASE-00": ("Foam insert, two layers", "1", "Closed-cell PE, 25 mm",
                    "BOTH layers now supplied as DXF at 1:1. Cut-outs labelled per RFQ M-06."),
    }
    seen = set()
    for row in t.rows[1:]:
        pid = row.cells[0].text.strip()
        if pid in updates:
            seen.add(pid)
            for i, v in enumerate(updates[pid], start=1):
                if i < len(row.cells):
                    row.cells[i].text = v
    for pid in ("HM-08", "HM-09"):
        if pid not in seen:
            r = t.add_row()
            r.cells[0].text = pid
            for i, v in enumerate(updates[pid], start=1):
                if i < len(r.cells):
                    r.cells[i].text = v
    for pid, vals in [("MP-01", ("Module mounting plate — NEW IN PACKAGE v2", "1",
                                 "PA12, MJF",
                                 "Carries every purchased module above the carrier; the "
                                 "modules connect by keyed 2.54 mm ribbon jumpers. "
                                 "ICD-EEG-006.")),
                      ("FIT-01", ("Fit-test coupon — NEW IN PACKAGE v2", "1 per batch",
                                  "PA12, same build as the batch",
                                  "Three bores at 9.20, 9.35 and 9.15 mm. Checked before the "
                                  "batch is accepted (QP-EEG-010)."))]:
        r = t.add_row()
        r.cells[0].text = pid
        for i, v in enumerate(vals, start=1):
            if i < len(r.cells):
                r.cells[i].text = v

    # ---- section 12 objection statuses ---------------------------------------------
    t = d.tables[11]
    status = {
        "1": " STATUS IN PACKAGE v2: unchanged and accepted.",
        "2": " STATUS IN PACKAGE v2: unchanged. Stage 0 must measure it.",
        "3": " STATUS IN PACKAGE v2: unchanged and still the most important open question. "
             "RFQ Rev D section 9.2 now names the measurement.",
        "4": " STATUS IN PACKAGE v2: unchanged. RFQ Rev D section 9.2 now requires recruitment "
             "across hair types and reporting by hair type rather than pooled.",
        "5": " STATUS IN PACKAGE v2: unchanged. The lights are dark at power-on as well as "
             "during recording, which is one more circumstance in which no photograph can "
             "show a lit device.",
        "6": " STATUS IN PACKAGE v2: unchanged. RFQ Rev D section 9.2 now requires the mass "
             "and centre-of-gravity measurement on the first prototype.",
        "7": " STATUS IN PACKAGE v2: unchanged and addressed in section 3.",
        "8": " STATUS IN PACKAGE v2: RFQ Rev D section 9.2 now requires twenty-five "
             "release-and-refit cycles with disinfectant exposure, and the FIT-01 coupon "
             "checks the fit on every batch.",
        "9": " STATUS IN PACKAGE v2: unchanged. The gland is a replaceable module.",
        "10": " STATUS IN PACKAGE v2: partly answered. The carrier, POD-P1, MP-01, HM-04, "
              "HM-08, HM-09, the pads and the coupon are now released as parametric STEP "
              "models with dimensioned drawings. HM-01 is still a rendered form study.",
        "11": " STATUS IN PACKAGE v2: unchanged and accepted.",
        "12": " STATUS IN PACKAGE v2: UNCHANGED AND STILL BLOCKING. RISK-EEG-011 is the pack "
              "the reviewer receives. No unit goes on a head before written sign-off.",
    }
    for row in t.rows[1:]:
        n = row.cells[0].text.strip()
        if n in status:
            c = row.cells[2]
            if c.paragraphs and c.paragraphs[-1].runs:
                r = c.paragraphs[-1].add_run(status[n])
                r.bold = True
            else:
                c.text = c.text + status[n]

    # ---- section 13 captions ---------------------------------------------------------
    for p in d.paragraphs:
        t0 = p.text.strip()
        if t0.startswith("FIG-22"):
            set_text(p, "FIG-22 — envelope detector. Superseded in Rev E by SCH-EEG-005 "
                        "sheet 3: a precision full-wave rectifier, a second-order Butterworth "
                        "low-pass at f0 = 48.8 Hz with Q = 0.74, and a buffered ×0.0909 "
                        "output, on one OPA4376 quad per channel. The Rev A dual op-amp left "
                        "the filter out of circuit (ECO-EEG-004) and its equal-R equal-C "
                        "network would have put the corner at 31 Hz (ECO-EEG-006).",
                     italic=True)
        elif t0.startswith("FIG-23"):
            set_text(p, "FIG-23 — one of sixteen identical input protection paths. In "
                        "package v2 these are on the carrier EEG-CAR-01 itself as R1–R16, "
                        "D1–D16 and C1–C16, not on a separate panel board. Clamp "
                        "leakage must stay below 1 nA or the lead-off reading — and "
                        "therefore the contact light on the head — is wrong. The current "
                        "budget is in RISK-EEG-011 section 4.", italic=True)
        elif t0.startswith("FIG-12") or t0.startswith("POD-00"):
            set_text(p, "FIG-12 — the pod as drawn for a consolidated board "
                        "(96 × 72 × 34 mm). SUPERSEDED: Phase 1 uses POD-P1, now a "
                        "released parametric model with the panel openings of RFQ M-02, and "
                        "Phases 2–3 use the 116 × 46 × 88 mm occipital shell on "
                        "HM-01. Note that the data and charge USB-C connectors are separate "
                        "(ECO-EEG-003).", italic=True)

    d.save(DST)
    print(f"wrote {DST}")
    print(f"  {renamed} figure labels renamed to FIG-nn")
    return DST


if __name__ == "__main__":
    main()
